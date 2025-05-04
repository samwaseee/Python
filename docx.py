from docx import Document
from docx.shared import Inches

def create_docx():
    # Create a new Document
    doc = Document()

    # Title
    doc.add_heading('Data Communication Answers', 0)

    # Sp22 Section
    doc.add_heading('Spring 2022', level=1)

    # Sp22 1.a
    doc.add_heading('1. a) Communication Between Computer R and S Across Hops', level=2)
    doc.add_paragraph(
        "The communication between computers R (port address k) and S involves multiple hops. Based on the TCP/IP model and encapsulation/decapsulation, the content of the transport, network, and data link layers for each hop is as follows:"
    )
    doc.add_paragraph(
        "- **Source Host (Computer R):**"
        "\n  - **Transport Layer:** Encapsulates the application message into a segment (TCP) or user datagram (UDP) with source port (k), destination port, and control information."
        "\n  - **Network Layer:** Encapsulates the segment into a datagram with source IP (R), destination IP (S), and routing information."
        "\n  - **Data Link Layer:** Encapsulates the datagram into a frame with source MAC (R) and destination MAC (next hop, e.g., router)."
    )
    doc.add_paragraph(
        "- **Router (Intermediate Hop):**"
        "\n  - **Data Link Layer (Incoming):** Decapsulates the frame to extract the datagram, verifying MAC addresses."
        "\n  - **Network Layer:** Inspects the datagram’s IP addresses and forwards it to the next hop."
        "\n  - **Data Link Layer (Outgoing):** Encapsulates the datagram into a new frame with the router’s MAC as source and the next hop’s MAC as destination."
    )
    doc.add_paragraph(
        "- **Destination Host (Computer S):**"
        "\n  - **Data Link Layer:** Decapsulates the frame to extract the datagram."
        "\n  - **Network Layer:** Decapsulates the datagram to extract the segment."
        "\n  - **Transport Layer:** Decapsulates the segment to deliver the message to the correct application process."
    )

    # Sp22 1.b
    doc.add_heading('1. b) Functions of ISO-OSI Model Layers', level=2)
    doc.add_paragraph(
        "The ISO-OSI model has seven layers, each with specific functions:"
        "\n1. **Physical Layer:** Transmits raw bits as electrical/optical signals, defines medium characteristics."
        "\n2. **Data Link Layer:** Ensures reliable frame transfer, handles error detection/correction, uses MAC addresses."
        "\n3. **Network Layer:** Routes packets using IP addresses, manages fragmentation and forwarding."
        "\n4. **Transport Layer:** Provides reliable process-to-process delivery, flow/error/congestion control, uses port numbers."
        "\n5. **Session Layer:** Manages communication sessions, handles recovery and synchronization."
        "\n6. **Presentation Layer:** Translates, encrypts, and compresses data for application compatibility."
        "\n7. **Application Layer:** Provides network services to applications (e.g., HTTP, SMTP)."
    )

    # Sp22 2.a
    doc.add_heading('2. a) Time and Frequency Domain Plots and Bandwidth', level=2)
    doc.add_paragraph(
        "- **Sine Wave:**"
        "\n  - **Time Domain:** Smooth, periodic oscillation (e.g., 5V peak, 8 Hz, oscillates between +5V and -5V every 1/8 s)."
        "\n  - **Frequency Domain:** Single spike at 8 Hz with 5V amplitude."
        "\n  - **Bandwidth:** 0 Hz (single frequency)."
        "\n- **Voice Signal:**"
        "\n  - **Time Domain:** Irregular, non-periodic waveform representing speech."
        "\n  - **Frequency Domain:** Continuous curve from 0 to 4 kHz."
        "\n  - **Bandwidth:** 4 kHz (4000 Hz)."
    )

    # Sp22 2.b
    doc.add_heading('2. b) Signal Power After Transmission', level=2)
    doc.add_paragraph(
        "Given: Power at end (P₂) = 3 mW, loss = 2 dB over 3 km."
        "\nUsing the decibel formula: -2 = 10 log₁₀(P₂/P₁), P₁ = P₂ / 10^(-2/10) ≈ 3 / 0.630957 ≈ 4.76 mW."
        "\n**Answer:** Initial power was 4.76 mW."
    )

    # Sp22 2.b Alternative
    doc.add_heading('2. b) Nyquist vs. Shannon Formulas', level=2)
    doc.add_paragraph(
        "- **Nyquist Bit Rate (Noiseless):** Bit Rate = 2 × Bandwidth × log₂(L)"
        "\n  - Assumes no noise, depends on bandwidth and signal levels (L)."
        "\n  - Example: 4 kHz, 4 levels → 16 kbps."
        "\n- **Shannon Capacity (Noisy):** C = Bandwidth × log₂(1 + SNR)"
        "\n  - Accounts for noise, depends on bandwidth and SNR."
        "\n  - Example: 4 kHz, SNR = 1000 → 40 kbps."
        "\n- **Comparison:** Nyquist is idealistic (no noise); Shannon is realistic (includes noise)."
        "\n- **Evaluation:** Nyquist sets an upper bound; Shannon guides practical design."
    )

    # Sp22 3.a
    doc.add_heading('3. a) Synchronous vs. Asynchronous Transmission and Signal Propagation', level=2)
    doc.add_paragraph(
        "**Synchronous vs. Asynchronous Transmission:**"
        "\n- **Synchronous:** Continuous data stream, shared clock, efficient for large transfers, complex."
        "\n- **Asynchronous:** Small data blocks, start/stop bits, simpler, less efficient."
        "\n**Signal Propagation:**"
        "\n- **Guided Media (e.g., Twisted-Pair, Coaxial, Fiber):** Signals travel through physical paths, face attenuation, distortion, noise. Fiber has lowest loss."
        "\n- **Unguided Media (e.g., Radio, Microwaves):** Signals travel wirelessly, face path loss, multipath fading, interference."
    )

    # Sp22 3.b Alternative
    doc.add_heading('3. b) Line Coding for Bit Stream 100100000011', level=2)
    doc.add_paragraph(
        "For bit stream 100100000011:"
        "\n- **NRZ-I:** Transitions for 1s (e.g., low → high for first 1, high → low for second 1)."
        "\n- **Manchester:** Mid-bit transitions (1: high-to-low, 0: low-to-high)."
        "\n- **Differential Manchester:** Mid-bit transition always, start-bit transition for 0s."
        "\n- **AMI:** 1s alternate +V/-V, 0s are 0V."
        "\n- **HDB3:** Like AMI, inserts violation pulse for four consecutive 0s."
    )

    # Au22 Section
    doc.add_heading('Autumn 2022', level=1)

    # Au22 1.a
    doc.add_heading('1. a) Message Size Calculation', level=2)
    doc.add_paragraph(
        "Given: Transmission time = 2 ms, data rate = 5 Mbps."
        "\nMessage Size = 0.002 s × 5 × 10⁶ bps = 10,000 bits = 1250 bytes."
        "\n**Answer:** Message size is 10,000 bits or 1250 bytes."
    )

    # Au22 1.b
    doc.add_heading('1. b) Topologies and Line Configuration', level=2)
    doc.add_paragraph(
        "- **Mesh:** Point-to-Point (dedicated links)."
        "\n- **Star:** Point-to-Point (links to hub)."
        "\n- **Bus:** Multipoint (shared backbone)."
        "\n- **Ring:** Point-to-Point (links to neighbors)."
    )

    # Au22 1.c
    doc.add_heading('1. c) Addressing in TCP/IP and Transport Layer Functions', level=2)
    doc.add_paragraph(
        "**Addressing:**"
        "\n- Required at application (names), transport (ports), network (IP), and data link (MAC) layers for accurate delivery."
        "\n**Transport Layer Functions:**"
        "\n- Encapsulation, connection management, flow control, error control, congestion control, process-to-process delivery."
    )

    # Au22 1.d
    doc.add_heading('1. d) Frequency Domain and Non-Periodic Signal', level=2)
    doc.add_paragraph(
        "**Periodic Signal:**"
        "\n- Bandwidth = 200 kHz, middle frequency = 140 kHz → Range: 40 kHz to 240 kHz."
        "\n- Frequency domain: Discrete spikes, highest amplitudes at 40 kHz and 240 kHz."
        "\n**Non-Periodic Signal:**"
        "\n- Bandwidth = 5 kHz → Infinitely many frequencies (continuous spectrum)."
    )

    # Au22 2
    doc.add_heading('2. Shannon Capacity and SNR', level=2)
    doc.add_paragraph(
        "Given: Capacity = 24 Mbps, bandwidth = 2 MHz."
        "\nSNR = 4095 (from 24 × 10⁶ = 2 × 10⁶ × log₂(1 + SNR))."
        "\nSNR_dB = 10 log₁₀(4095) ≈ 36.12 dB."
        "\n**Answer:** SNR = 4095, SNR_dB = 36.12 dB."
    )

    # Au22 2 Alternative
    doc.add_heading('2. a) Signal Power After Amplification', level=2)
    doc.add_paragraph(
        "Given: Power at Y = P_X/3, amplified by 5x."
        "\nLoss (X to Y): -4.77 dB."
        "\nGain: 6.99 dB."
        "\nOverall: -4.77 + 6.99 = 2.22 dB (gain)."
        "\n**Answer:** Overall gain is 2.22 dB."
    )

    # Au22 2.b
    doc.add_heading('2. b) Digital Signal and Line Coding', level=2)
    doc.add_paragraph(
        "**i) Bit Rate 4 bps, 1 Bit/Element:** Square wave, 0.25 s/bit, 2 levels (e.g., 1010 → high-low-high-low)."
        "\n**ii) Bit Rate 8 bps, 2 Bits/Element:** 4 levels, 0.25 s/element (e.g., 11001001 → 3V-0V-2V-1V)."
        "\n**iii) Line Coding:** Converts bits to signals (e.g., NRZ-L, Manchester) for synchronization, error minimization."
    )

    # Au22 3.a
    doc.add_heading('3. a) PCM for Sinusoidal Signal', level=2)
    doc.add_paragraph(
        "Given: Sinusoid (-16V to +16V), 4 bits/sample."
        "\n- Quantization: 16 levels, 2V steps (-16V to 16V)."
        "\n- Encoding: e.g., -16V → 0000, 6V → 1011."
        "\nExample: 5.7V → 6V (1011), -9.2V → -10V (0011)."
    )

    # Au22 3.b Alternative
    doc.add_heading('3. b) Digitizing Human Voice', level=2)
    doc.add_paragraph(
        "Given: Voice bandwidth = 104 kHz, 1 bit/sample."
        "\n- Sampling Rate: 2 × 104 kHz = 208 kHz."
        "\n- Bit Rate: 208 kHz × 1 = 208 kbps."
        "\n- Bandwidth Ratio: Digital (104 kHz) / Analog (104 kHz) = 1."
    )

    # Sp23 Section
    doc.add_heading('Spring 2023', level=1)

    # Sp23 1.a
    doc.add_heading('1. a) Fill in the Gaps', level=2)
    doc.add_paragraph(
        "i. **application** process delivers the entire message."
        "\nii. **physical** layer is closest to the medium."
        "\niii. **physical** layer converts bits to signals."
        "\niv. **logical address** changes over the Internet."
        "\nv. Internet model ensures **reliability**."
    )

    # Sp23 1.b
    doc.add_heading('1. b) ISO-OSI Layer Functions', level=2)
    doc.add_paragraph(
        "(Same as Sp22 1.b)"
    )

    # Sp23 2.a
    doc.add_heading('2. a) Transmission and Propagation Time, Delay-Bandwidth Product', level=2)
    doc.add_paragraph(
        "Given: Message = 500 Mbyte, bandwidth = 2 Gbps, distance = 10,000 km, speed = 2.4 × 10⁸ m/s."
        "\n- Transmission Time: (4 × 10⁹ bits) / (2 × 10⁹ bps) = 2 s."
        "\n- Propagation Time: 10⁷ m / (2.4 × 10⁸ m/s) ≈ 41.67 ms."
        "\n- Delay-Bandwidth Product: 0.04167 s × 2 × 10⁹ bps ≈ 83.34 Mbits."
    )

    # Sp23 2
    doc.add_heading('2. TV Channel Data Rate and Levels', level=2)
    doc.add_paragraph(
        "Given: Bandwidth = 500 MHz, SNR_dB = 15.2."
        "\n- Capacity: 500 × 10⁶ × log₂(1 + 33.113) ≈ 2.545 Gbps."
        "\n- Levels: L ≈ 4 (from 2 × 500 × 10⁶ × log₂(4) = 2 Gbps, feasible)."
    )

    # Sp23 3.a
    doc.add_heading('3. a) Transmission Impairments and Mitigation', level=2)
    doc.add_paragraph(
        "- **Impairments:** Attenuation (energy loss), distortion (shape change), noise (interference)."
        "\n- **Mitigation:** Amplifiers, repeaters, equalizers, shielding, error correction."
    )

    # Sp23 3.b
    doc.add_heading('3. b) Signal Power After Cable Loss and Amplification', level=2)
    doc.add_paragraph(
        "Given: P₁ = y mW, loss = -1 dB/km, path: 4 km → -3 dB → 3 km → +3 dB → 1 km."
        "\nTotal dB = -4 - 3 - 3 + 3 - 1 = -8 dB."
        "\nP₂ = 0.1585 × y mW."
    )

    # Sp23 3.b Alternative
    doc.add_heading('3. b) Fiber Optic Signal Transmission', level=2)
    doc.add_paragraph(
        "Fiber optics carry light via total internal reflection in a glass/plastic core, encoding data as pulses."
        "\n- High bandwidth, low attenuation, noise immunity."
    )

    # Sp24 Section
    doc.add_heading('Spring 2024', level=1)

    # Sp24 1.a
    doc.add_heading('1. a) Network Topologies: Advantages and Disadvantages', level=2)
    doc.add_paragraph(
        "- **Mesh:** Robust, reliable; expensive, complex."
        "\n- **Star:** Cost-effective, easy to manage; hub is single point of failure."
        "\n- **Bus:** Simple, cost-effective; backbone failure disrupts network."
        "\n- **Ring:** Easy to install; ring break can disable network."
    )

    # Sp24 1.b
    doc.add_heading('1. b) Components of Data Communication System', level=2)
    doc.add_paragraph(
        "- **Message:** Data to communicate."
        "\n- **Sender:** Device sending the message."
        "\n- **Receiver:** Device receiving the message."
        "\n- **Medium:** Physical path (e.g., cables, wireless)."
        "\n- **Protocol:** Rules for communication."
        "\nFigure: Sender → Medium → Receiver, coordinated by protocol."
    )

    # Sp24 2.a
    doc.add_heading('2. a) Simplifying Data Communication', level=2)
    doc.add_paragraph(
        "- **Protocol Layering:** Divides tasks into layers (e.g., TCP/IP)."
        "\n- **Standardized Protocols:** Ensure interoperability."
        "\n- **Encapsulation:** Simplifies data handling."
        "\n- **Hierarchical Structure:** Abstracts lower-layer complexity."
    )

    # Sp24 2.b
    doc.add_heading('2. b) Transmission Impairments and Bit Rate', level=2)
    doc.add_paragraph(
        "(Impairments same as Sp23 3.a)"
        "\nGiven: SNR_dB = 36, bandwidth = 1 MHz."
        "\n- Bit Rate: 1 × 10⁶ × log₂(1 + 3981.07) ≈ 11.96 Mbps."
        "\n- Signal Levels: L ≈ 64."
    )

    # Sp24 3.a
    doc.add_heading('3. a) Advantages of Optical Fiber', level=2)
    doc.add_paragraph(
        "- Higher bandwidth, lower attenuation, noise immunity, security, smaller size, safety, long-distance transmission."
    )

    # Sp24 3.b
    doc.add_heading('3. b) NRZ-I and Manchester for Bit Stream 11010011', level=2)
    doc.add_paragraph(
        "- **NRZ-I:** Transitions for 1s (e.g., low → high for first 1, high → low for second 1)."
        "\n- **Manchester:** Mid-bit transitions (1: high-to-low, 0: low-to-high)."
    )

    # Save the document
    doc.save('DataCommunicationAnswers.docx')

if __name__ == '__main__':
    create_docx()